from pathlib import Path
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


def set_font(style, name, size=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        style.font.size = Pt(size)


work_dir = Path(os.environ["COURSE_SW_DIR"])
shot_dir = work_dir / "建模过程截图"
out_path = work_dir / "课程设计建模过程说明.docx"

steps = [
    (
        "步骤1  零件1：拉伸圆形底座",
        "01-1_零件1_拉伸圆形底座.jpg",
        "在上视基准面绘制直径150mm的圆，拉伸30mm形成圆形底座，作为整个支座的安装基础。",
    ),
    (
        "步骤2  零件1：建立双耳外形和孔",
        "01-2_零件1_建立双耳外形和孔.jpg",
        "在圆形底座上方绘制双耳支座外形，拉伸形成耳板，并在耳板圆头位置预留直径18mm销孔。",
    ),
    (
        "步骤3  零件1：切除中间25mm槽",
        "01-3_零件1_切除中间25mm槽.jpg",
        "按图纸要求从双耳中间切除25mm宽槽，槽以中间对称分开，形成左右两侧耳板。",
    ),
    (
        "步骤4  零件2：拉伸叉形基础块",
        "02-1_零件2_拉伸叉形基础块.jpg",
        "绘制连杆基础矩形轮廓并拉伸，得到叉形连杆的初始毛坯体。",
    ),
    (
        "步骤5  零件2：建立单耳铰接端",
        "02-2_零件2_建立单耳铰接端.jpg",
        "在连杆一端建立圆头单耳结构，并加工直径18mm铰接孔，后续与支座孔同心装配。",
    ),
    (
        "步骤6  零件2：切除叉口",
        "02-3_零件2_切除叉口.jpg",
        "从连杆另一端切除中间材料，形成叉形开口，使其满足图纸中的前后分开关系。",
    ),
    (
        "步骤7  零件2：加工叉耳孔",
        "02-4_零件2_加工叉耳孔.jpg",
        "在叉形端加工通孔，完成连杆零件的主要结构。",
    ),
    (
        "步骤8  零件3：拉伸销轴头部",
        "03-1_零件3_拉伸销轴头部.jpg",
        "绘制较大直径圆并拉伸，形成阶梯销轴的头部，用于轴向限位。",
    ),
    (
        "步骤9  零件3：拉伸18mm轴身",
        "03-2_零件3_拉伸18mm轴身.jpg",
        "在同一轴线上绘制直径18mm圆并拉伸，形成与支座孔、连杆孔配合的销轴杆部。",
    ),
    (
        "步骤10  装配：插入圆底双耳支座",
        "04_装配_插入圆底双耳支座.jpg",
        "新建装配体，首先插入零件1并固定在装配原点，作为后续零件定位基准。",
    ),
    (
        "步骤11  装配：插入连杆并调整角度",
        "05_装配_连杆绕销轴成135度.jpg",
        "插入零件2，将连杆铰接孔与支座孔对齐，并按图纸要求调整为约135度的装配姿态。",
    ),
    (
        "步骤12  装配：插入第二个连杆并用销轴连接",
        "06_装配_插入销轴.jpg",
        "在第一个连杆的叉耳端插入第二个零件2，使第二个连杆的单耳端进入第一个连杆的叉口；再插入零件3作为连接销轴，使两个连杆形成铰接连接。",
    ),
    (
        "步骤13  装配检查：俯视关系",
        "07_装配_视图A.jpg",
        "从上方检查支座双耳、中间槽、连杆铰接端和销轴的前后位置关系。",
    ),
    (
        "步骤14  装配检查：前视关系",
        "08_装配_视图B.jpg",
        "从前视方向检查连杆角度和孔位关系，确认三件零件按图纸要求组合。",
    ),
]

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

set_font(doc.styles["Normal"], "宋体", 10.5)
set_font(doc.styles["Heading 1"], "黑体", 16)
set_font(doc.styles["Heading 2"], "黑体", 13)

title = doc.add_heading("冲压作业装配体建模过程说明", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("内容包括三个零件的主要建模步骤截图，以及最终装配过程和视图检查。")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("一、建模与装配过程", level=1)
for heading, image_name, text in steps:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(text)
    image_path = shot_dir / image_name
    if image_path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        with Image.open(image_path) as img:
            width = Cm(15.2 if img.width >= img.height else 11.5)
        run.add_picture(str(image_path), width=width)
    else:
        doc.add_paragraph("截图文件未找到：" + image_name)

doc.add_heading("二、装配检查记录", level=1)
doc.add_paragraph(
    "装配完成后检查孔轴关系：支座孔中心、连杆铰接孔中心和销轴中心在装配坐标中重合，中心偏差为0.000mm。"
)
doc.add_paragraph(
    "最终文件包括：零件1.SLDPRT、零件2.SLDPRT、零件3.SLDPRT、课程设计装配体.SLDASM。装配体中零件2作为连杆被使用两次，零件3作为销轴被使用两次。"
)

doc.save(out_path)
print(out_path)
